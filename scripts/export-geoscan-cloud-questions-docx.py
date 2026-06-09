from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/darnhalm/Documents/CURSOR/model-viewer/docs/api/GEOSCAN-CLOUD-QUESTIONS-RU.docx"


QUESTIONS = [
    {
        "area": "Авторизация и Geoscan ID",
        "question": "Подтвердить схему авторизации через Geoscan ID: OAuth/OIDC, пользовательский access token, service token для backend Heritage3D, editor session token для iframe-редактора и viewer/embed token для просмотра.",
        "why": "Нужно разделить действия пользователя, системные операции backend, встроенное редактирование и публичный просмотр.",
        "ask": "Зафиксировать типы токенов, TTL, scopes, allowed origins и правила аудита действий.",
    },
    {
        "area": "Лимиты и тарифы",
        "question": "Уточнить, могут ли тариф, хранилище, месячные лимиты и остатки быть привязаны не только к личному Geoscan ID / пользовательскому токену, но и к организации/музею.",
        "why": "Для корпоративной подписки логичен сценарий, где лимиты принадлежат рабочей группе или музею, а пользователи расходуют общий организационный лимит.",
        "ask": "Нужны поля и API для лимитов на уровне organizationId: тариф, общий объем, использовано, месячные лимиты обработки/загрузок, участники, доступные остатки для пользователя в рамках организации.",
    },
    {
        "area": "Справочники форматов",
        "question": "Подтвердить отдельный справочник поддерживаемых форматов файлов и расширений.",
        "why": "Geoscan Cloud может добавлять новые форматы: Cesium 3D Tiles, GLB/GLTF, 2D scan, panorama_360, point cloud, raster, vector и другие.",
        "ask": "Нужен endpoint, который портал сможет читать динамически без доработки frontend при появлении новых форматов.",
    },
    {
        "area": "CRS / системы координат",
        "question": "Подтвердить отдельный справочник CRS / систем координат с поиском по EPSG-коду или названию.",
        "why": "CRS не должен смешиваться со справочником форматов, потому что это отдельный шаг настройки загрузки и отдельный источник ошибок.",
        "ask": "Нужны поддерживаемые CRS, локальные координаты, единицы измерения и подсказки для UI при загрузке.",
    },
    {
        "area": "TW_ID и тип представления",
        "question": "Подтвердить, что TW_ID содержит тип культурного ассета / представления: cesium_tileset, glb_gltf, scan_2d, panorama_360 и т.д.",
        "why": "На сайте один H3D_ID может иметь несколько TW_ID разных типов: 3D-модель, 2D-скан, панорама, версия модели или другой медиаобъект.",
        "ask": "Нужна стабильная enum-схема representationType и связь TW_ID с geoscanProjectId, datasetId, geoscanObjectId, tilesetId, manifestUrl и previewUrl.",
    },
    {
        "area": "Встроенный редактор",
        "question": "Подтвердить возможность встроить Geoscan Cloud editor в портал Heritage3D через iframe/editor session.",
        "why": "Через REST API невозможно полноценно повторить интерфейс редактирования модели, POI, туров, слоев, измерений и preview.",
        "ask": "Нужен editorUrl/editorToken, список разрешенных инструментов, кнопка сохранения сцены и событие/endpoint публикации результата.",
    },
    {
        "area": "Публичный viewer и embed",
        "question": "Уточнить, когда публичный viewer может открываться без отдельного viewer token, а когда нужен короткоживущий embed token.",
        "why": "Для полностью публичных моделей можно упростить просмотр, но для черновиков, закрытых коллекций, модерации и запрета скачивания нужен контролируемый доступ.",
        "ask": "Нужны правила для public embed, restricted embed, manifestUrl, iframeUrl, previewUrl и запрета скачивания исходников.",
    },
    {
        "area": "Файлы, слои и пакетные операции",
        "question": "Подтвердить API для пакетного управления файлами, слоями и версиями внутри TW_ID.",
        "why": "На портале версии модели и разные даты съемки должны управляться как слои или связанные ассеты, без показа пользователю стандартной файловой структуры облака.",
        "ask": "Нужны операции показать/скрыть, удалить, подлететь, переименовать, переместить, сгруппировать, назначить порядок и связать с TW_ID_Layer.",
    },
    {
        "area": "Preview и миниатюры",
        "question": "Уточнить API для preview модели, preview точки POI и preview каталога/коллекции.",
        "why": "На сайте нужны обложки объектов, коллекций и публичных карточек, а в редакторе уже есть действие 'Сохранить ракурс'.",
        "ask": "Нужны previewUrl, дата обновления, автор, источник preview и возможность назначить preview для каталога/коллекции.",
    },
    {
        "area": "Webhook-и и синхронизация",
        "question": "Подтвердить webhook-и о статусе загрузки, обработки, публикации, ошибках, изменении preview и удалении данных.",
        "why": "Heritage3D должен обновлять статус ассета и публичной карточки без ручного опроса Geoscan API.",
        "ask": "Нужна подпись webhook, retry policy, eventId, timestamp, H3D_ID/TW_ID mapping и единая схема ошибок.",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.first_child_found_in("w:tblCellMar")
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tbl_cell_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT


def format_paragraph(paragraph, size=11, color="111827", bold=False):
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor.from_string(color)
        run.bold = bold


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.15

for style_name, size, color, before, after in [
    ("Heading 1", 20, "000000", 20, 6),
    ("Heading 2", 16, "000000", 18, 6),
    ("Heading 3", 14, "434343", 16, 4),
]:
    style = styles[style_name]
    style.font.name = "Arial"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = False
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15

title = doc.add_paragraph()
title_run = title.add_run("Вопросы к согласованию по Geoscan Cloud API")
title_run.font.name = "Arial"
title_run.font.size = Pt(26)
title_run.font.color.rgb = RGBColor(0, 0, 0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_after = Pt(3)

subtitle = doc.add_paragraph()
subtitle.add_run("Heritage3D / Geoscan Cloud. Рабочий список вопросов для финализации API и встраиваемого редактора.")
format_paragraph(subtitle, size=11, color="475569")

lead = doc.add_paragraph()
lead.add_run(
    "Документ фиксирует вопросы, которые нужно отдельно подтвердить с командой Geoscan. "
    "Основной акцент: авторизация через Geoscan ID, лимиты и тарифы, справочники форматов/CRS, "
    "связь H3D_ID и TW_ID, встроенный редактор, embed/viewer и синхронизация статусов."
)
format_paragraph(lead)

doc.add_heading("Ключевые вопросы", level=1)

table = doc.add_table(rows=1, cols=4)
set_table_width(table)
set_cell_margins(table)
table.style = "Table Grid"
headers = ["Область", "Вопрос", "Почему важно", "Что нужно уточнить"]
widths = [1500, 3200, 2500, 2160]

for idx, text in enumerate(headers):
    cell = table.rows[0].cells[idx]
    cell.text = text
    set_cell_shading(cell, "E5E7EB")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        format_paragraph(paragraph, size=10, color="111827", bold=True)

for item in QUESTIONS:
    row = table.add_row()
    values = [item["area"], item["question"], item["why"], item["ask"]]
    for idx, value in enumerate(values):
        cell = row.cells[idx]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        for paragraph in cell.paragraphs:
            format_paragraph(paragraph, size=9, color="111827", bold=False)

for row in table.rows:
    for idx, width in enumerate(widths):
        row.cells[idx].width = width

doc.add_heading("Отдельный акцент по лимитам", level=1)
paragraph = doc.add_paragraph()
paragraph.add_run("Важно не ограничиваться моделью, где лимиты привязаны только к личному пользователю Geoscan ID. ")
paragraph.add_run(
    "Для музеев и организаций нужен корпоративный сценарий: тариф и общий лимит хранилища/обработки закреплены за organizationId, "
    "а конкретные пользователи работают внутри этой организации по своим ролям и расходуют общий организационный лимит."
)
format_paragraph(paragraph)

paragraph = doc.add_paragraph()
paragraph.add_run(
    "В API желательно явно различить: личный пользовательский контекст, доступные пользователю организации, лимиты выбранной организации, "
    "остаток лимита для текущего действия и причины отказа при превышении квоты."
)
format_paragraph(paragraph)

doc.save(OUTPUT)
print(OUTPUT)
