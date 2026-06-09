from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "/Users/darnhalm/Documents/CURSOR/model-viewer/docs/api/Geoscan Cloud API для Heritage3D.docx"


def set_run(run, size=11, color="111827", bold=False, italic=False, font="Calibri"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_text(text, *, bold_prefix=None, italic=False, after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        set_run(paragraph.add_run(bold_prefix), bold=True)
        set_run(paragraph.add_run(text[len(bold_prefix):]), italic=italic)
    else:
        set_run(paragraph.add_run(text), italic=italic)
    return paragraph


def add_bullets(items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run(paragraph.add_run(item))


def add_numbered(items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.line_spacing = 1.25
        set_run(paragraph.add_run(item))


def add_code(lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.22)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.1
    for index, line in enumerate(lines.splitlines()):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run(run, size=9.5, color="334155", font="Consolas")
    return paragraph


def add_note(label, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.08)
    paragraph.paragraph_format.line_spacing = 1.2
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    set_run(paragraph.add_run(f"{label}: "), bold=True, color="1F3A5F")
    set_run(paragraph.add_run(text), color="334155")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Страница ")
    set_run(run, size=9, color="64748B")
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.85)
section.right_margin = Inches(0.9)
section.bottom_margin = Inches(0.85)
section.left_margin = Inches(0.9)
section.header_distance = Inches(0.42)
section.footer_distance = Inches(0.42)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for style_name, size, color, before, after in [
    ("Heading 1", 16, "2E74B5", 18, 10),
    ("Heading 2", 13, "2E74B5", 14, 7),
    ("Heading 3", 12, "1F4D78", 10, 5),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = True
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing = 1.15

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_run(header.add_run("Heritage3D | Интеграция с Geoscan Cloud"), size=9, color="64748B")
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(8)
title.paragraph_format.space_after = Pt(3)
set_run(title.add_run("Geoscan Cloud API для Heritage3D"), size=24, color="0B2545", bold=True)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(14)
set_run(
    subtitle.add_run(
        "Требования к API, встраиваемому редактору и защищенному доступу к данным"
    ),
    size=12,
    color="475569",
    italic=True,
)

add_note(
    "Статус документа",
    "Рабочий перечень требований для обсуждения с командой Geoscan Cloud. "
    "Указанные в Excel-таблице URL являются условными эндпойнтами по образцу Sketchfab "
    "и должны быть уточнены с разработчиками Geoscan Cloud.",
)

add_text(
    "Цель интеграции — использовать Geoscan Cloud как техническую платформу хранения, "
    "обработки, редактирования и просмотра цифровых моделей, сохранив на стороне "
    "heritage3d.ru карточки культурных объектов, метаданные, каталоги, модерацию и публичную обвязку."
)

doc.add_heading("1. Geoscan Cloud Embedded Editor", level=1)
add_text(
    "Нам нужна встройка окна редактирования модели из Geoscan Cloud внутрь портала Heritage3D. "
    "По смыслу это аналог панели управления моделью на Sketchfab: пользователь редактирует сцену, "
    "слои, туры, POI, preview и настройки отображения прямо на портале, но фактически работает "
    "внутри встроенного редактора Geoscan Cloud."
)
add_text(
    "Поскольку авторизация пользователей предполагается через Geoscan ID, желательно обеспечить "
    "бесшовный вход во встроенный редактор. Пользователей можно объединять в команды Geoscan Cloud. "
    "На Heritage3D такие команды могут соответствовать организациям: музеям, исследовательским "
    "центрам и другим учреждениям."
)
add_text(
    "Внутренние роли портала Heritage3D — администратор, модератор, редактор и пользователь — "
    "остаются на стороне портала и регулируют доступ к карточкам объектов, модерации и публикации."
)
add_text(
    "Отдельно действуют роли команды Geoscan Cloud: Владелец, Администратор, Участник и "
    "Внешний пользователь. Они определяют членство в команде, доступ к общему диску и базовые "
    "возможности работы с облачными данными."
)
add_text(
    "Внутри каталога Geoscan Cloud дополнительно используются роли Менеджер, Редактор, Аналитик, "
    "Зритель и Без доступа. Они назначаются для папки, объекта или проекта и наследуются вниз по "
    "иерархии. Для Внешнего пользователя роль каталога не может быть выше роли Аналитик."
)
add_text(
    "Роли Geoscan Cloud и роли Heritage3D не нужно автоматически приравнивать друг к другу: "
    "Geoscan Cloud контролирует доступ к файлам, проектам и редактору, а Heritage3D — доступ к "
    "карточкам культурных объектов, модерации и публикации на портале. При необходимости портал "
    "может хранить отдельное сопоставление ролей для пользовательских сценариев, но итоговые права "
    "на облачные ресурсы должны проверяться Geoscan Cloud по пользовательскому токену."
)
add_text(
    "Справка Geoscan Cloud: «Как управлять командой» — "
    "https://geoscan.helpdeskeddy.com/ru/knowledge_base/art/370/cat/181/kak-upravljat-komandoj"
)
add_text(
    "Отдельно используются роли команды Geoscan Cloud, которые определяют доступ пользователя к данным "
    "и функциям самого облака: Владелец, Администратор, Участник и Внешний пользователь. Владелец имеет "
    "полные права в рамках аккаунта. Администратор может управлять диском, данными и пользователями команды, "
    "но не безопасностью и брендированием. Участник и Внешний пользователь получают доступ к данным только "
    "после назначения роли в конкретном каталоге."
)
add_text(
    "В Geoscan Cloud также есть роли каталога: Менеджер, Редактор, Аналитик, Зритель и Без доступа. "
    "Они назначаются для папки, объекта или проекта и наследуются вниз по иерархии. Внешнему пользователю "
    "нельзя назначить роль каталога выше Аналитика. Эти роли регулируют работу с облачными файлами, проектами "
    "и слоями и не должны автоматически приравниваться к ролям Heritage3D. Портал хранит собственную роль "
    "пользователя, а при обращении к Geoscan Cloud дополнительно учитывает его команду и effective permissions "
    "для конкретного каталога или проекта."
)
add_text(
    "Справка Geoscan Cloud о командах и ролях: "
    "https://geoscan.helpdeskeddy.com/ru/knowledge_base/art/370/cat/181/kak-upravljat-komandoj",
    italic=True,
)
add_text(
    "Пользователь предварительно входит на Heritage3D через Geoscan ID, после чего портал получает "
    "пользовательский geoscan_user_access_token. При создании новой модели Heritage3D сначала создает "
    "черновик карточки модели и TW_ID на стороне портала."
)
add_text(
    "Для черновика backend определяет parentFolderId папки Geoscan Cloud, связанной с выбранным "
    "каталогом или коллекцией Heritage3D, и от имени пользователя создает пустой проект Geoscan Cloud. "
    "Полученный geoscanProjectId сохраняется в связи с TW_ID. Затем также от имени пользователя портал "
    "запрашивает короткоживущую editor session для данного проекта и открывает editorUrl во встроенном iframe. "
    "Пользователь загружает файлы, добавляет уже загруженные в облако данные, например модели из Agisoft, "
    "и редактирует сцену внутри проекта. До публикации модель остается черновиком на стороне Heritage3D."
)

doc.add_heading("2. Структура хранения и S3", level=1)
add_text(
    "Нужно вынести в API структуру хранения Geoscan Cloud: папки, проекты, "
    "наборы данных, слои и связанные сущности."
)
add_text(
    "Geoscan Cloud использует S3-хранилище для физических файлов. Портал Heritage3D не получает "
    "прямой доступ к bucket, object key или учетным данным S3. Все операции с файлами выполняются "
    "через API Geoscan Cloud: загрузка, обработка, публикация, перемещение, удаление и скачивание. "
    "Geoscan Cloud самостоятельно управляет связью между проектами, слоями и физическими объектами в S3."
)
add_text(
    "На портале мы не хотим напрямую показывать стандартную файловую структуру облака. Ее нужно "
    "обернуть в интерфейс Heritage3D: коллекции, каталоги, карточки культурных объектов и разделы "
    "моделей, используя стандартную структуру хранения Geoscan Cloud на backend-уровне."
)
add_code("collectionId Heritage3D → folderId Geoscan Cloud")
add_text(
    "Папка Geoscan Cloud может отображаться на Heritage3D как каталог или коллекция. "
    "При создании проекта backend передает parentFolderId папки, связанной с выбранной коллекцией. "
    "collectionId остается внутренним идентификатором Heritage3D."
)
add_text(
    "Для версий одной цифровой модели предпочтительно использовать слои внутри одного проекта: "
    "разные даты съемки, обработки и состояния до или после реставрации."
)

doc.add_heading("3. Связка H3D_ID, TW_ID и проекта Geoscan Cloud", level=1)
add_text(
    "На стороне Heritage3D есть сквозной идентификатор культурного объекта — H3D_ID. "
    "К нему может быть привязан один или несколько цифровых двойников или медиаобъектов TW_ID."
)
add_code(
    "H3D_ID\n"
    "  ├── TW_ID 1: 3D-модель\n"
    "  ├── TW_ID 2: 2D-скан\n"
    "  ├── TW_ID 3: 360-панорама\n"
    "  └── TW_ID 4: другое цифровое представление"
)
add_text(
    "TW_ID создается и хранится на стороне Heritage3D. Основная техническая связь:"
)
add_code(
    "H3D_ID\n"
    "  └── TW_ID\n"
    "       └── geoscanProjectId\n"
    "            ├── datasetId: слой 1\n"
    "            ├── datasetId: слой 2\n"
    "            └── datasetId: слой 3\n"
    "                 └── S3 object managed by Geoscan Cloud"
)
add_text(
    "Если API Geoscan Cloud поддерживает внешние идентификаторы, можно предусмотреть поле externalId, "
    "но источником TW_ID остается портал Heritage3D. TW_ID связывается с проектом Geoscan Cloud, "
    "а не с S3-файлом. Это позволяет Geoscan Cloud менять внутреннюю структуру хранения, CDN и правила "
    "доступа без доработки Heritage3D."
)
add_text(
    "Метаинформация культурного объекта — Dublin Core, CIDOC CRM, описания, связи с каталогами, "
    "статусы модерации и публикации — остается на стороне Heritage3D."
)

doc.add_heading("4. Управление файлами и слоями", level=1)
add_text("На уровне API требуется пообъектное и пакетное управление файлами и слоями:")
add_bullets(
    [
        "создание проекта и привязка к папке Geoscan Cloud;",
        "загрузка данных через embedded editor;",
        "переименование, перемещение и копирование;",
        "показ и скрытие слоев;",
        "удаление;",
        "получение статуса обработки;",
        "получение preview объекта и папки;",
        "получение временной ссылки на скачивание, если скачивание разрешено.",
    ]
)
add_text(
    "Если загрузка выполняется не через встроенный редактор, а непосредственно с портала, API Geoscan "
    "Cloud может выдавать короткоживущую signed upload URL для загрузки файла в S3. После загрузки портал "
    "передает идентификатор upload-сессии Geoscan Cloud для запуска обработки. Для обычного пользовательского "
    "сценария предпочтительна загрузка через embedded editor."
)
add_text("Для слоев нужны отдельные API-команды:")
add_bullets(
    [
        "получить список слоев и свойства конкретного слоя;",
        "показать или скрыть слой;",
        "удалить слой;",
        "изменить порядок;",
        "перенести или скопировать слой в другой проект;",
        "получить формат, размер, CRS, дату версии и статус обработки;",
        "получить временную ссылку на скачивание.",
    ]
)
add_text("Пакетные операции целесообразно разделить по уровням:")
add_bullets(
    [
        "операции над диском организации или команды;",
        "операции внутри папки, отображаемой как каталог или коллекция;",
        "операции над проектами внутри выбранной папки или коллекции;",
        "операции над слоями внутри проекта;",
        "системная операция backend Heritage3D для смешанного набора ресурсов.",
    ]
)
add_text(
    "Если пользователь выбирает объекты по TW_ID, backend Heritage3D сначала преобразует их в конкретные "
    "ID ресурсов Geoscan Cloud, после чего отправляет запрос на массовую операцию."
)

doc.add_heading("5. POI, туры и измерения", level=1)
add_text(
    "Нужно вынести в API POI, туры, измерения и связанные настройки внутри проекта, связанного с TW_ID."
)
add_text("Для POI нужны следующие параметры:")
add_bullets(
    [
        "идентификатор, название и описание;",
        "Markdown-разметка для форматированного текста, ссылок и изображений;",
        "цвет, позиция, камера и preview;",
        "длительность перехода и порядок внутри тура;",
        "связанная подсвечиваемая область;",
        "связанное измерение.",
    ]
)
add_text(
    "Важно, чтобы данными можно было управлять со страницы портала. Например, пользователь нажимает "
    "на дату, термин или точку интереса в статье, а встроенный плеер синхронно переходит к нужной POI, "
    "области модели или ракурсу."
)
add_text(
    "Команды и события viewer нужны для связи меток, ссылок и фрагментов текста на странице сайта "
    "с управлением интерактивным содержимым в плеере: турами, POI, объектами, измерениями и сравнением. "
    "Перечень предлагаемых команд и событий приведен на листе «Embed и JS» в таблице API."
)
add_text(
    "Желательно добавить возможность выделять не только точку, но и область модели в виде многоугольника. "
    "Это позволит размечать фрагменты объекта и связывать их с текстом статьи, POI или измерениями."
)
add_text("Для измерений нужны:")
add_bullets(
    [
        "единицы измерения;",
        "расстояния, площади и объемы;",
        "связанные подписи;",
        "экспорт в GeoJSON при необходимости.",
    ]
)
add_text(
    "Через API также нужно получать Bounding box сцены: примерные физические размеры модели по X/Y/Z "
    "и единицу измерения. Это позволит использовать габариты модели в карточке ассета, каталоге и связанных материалах."
)
add_text(
    "Отдельно требуется обсудить возможность калибровки масштаба модели с портала по известному расстоянию "
    "между двумя точками. После калибровки Geoscan Cloud должен возвращать обновленные единицы измерения, "
    "коэффициент масштаба и Bounding box."
)

doc.add_heading("6. Preview, сохранение сцены и публикация", level=1)
add_text(
    "Через API должна быть доступна preview-картинка модели, созданная через кнопку «Сохранить вид». "
    "Она используется на портале в карточке модели, каталоге и на публичной странице."
)
add_text(
    "Также желательно предусмотреть API для preview папок Geoscan Cloud или возможность получать набор "
    "preview моделей внутри папки. Это позволит назначать изображения-заставки для каталогов, коллекций и разделов."
)
add_text(
    "После редактирования модели во встроенном редакторе нужна явная кнопка «Сохранить», чтобы у пользователя "
    "не было путаницы относительно текущего состояния проекта. Сохранение должно фиксировать состояние и порядок "
    "слоев, видимость, камеру, preview, туры, POI, измерения и настройки viewer. Публикация или обновление "
    "публичного manifest выполняется отдельным действием."
)

doc.add_heading("7. Статусы, ошибки и события", level=1)
add_text(
    "Через API и webhook нужно получать состояния обработки, публикации и ошибки сцены. Heritage3D должен "
    "сохранять события в журнал и показывать пользователю понятные уведомления, включая ошибки нехватки ресурсов "
    "при публикации, например Insufficient resources to complete publication."
)
add_bullets(
    [
        "текущее состояние сцены;",
        "этап обработки и прогресс;",
        "статус публикации;",
        "код и текст ошибки;",
        "дата события;",
        "рекомендации для пользователя;",
        "журнал событий проекта;",
        "webhook о готовности, ошибке обработки или публикации.",
    ]
)
add_text(
    "Webhook нужен для автоматического уведомления пользователя с портала о завершении обработки и загрузки файла."
)

doc.add_heading("8. Тарифы и лимиты", level=1)
add_text("Через API нужно получать состояние лицевого счета и лимитов Geoscan Cloud:")
add_bullets(
    [
        "текущий тариф;",
        "доступное и занятое место;",
        "месячные лимиты;",
        "лимиты обработки;",
        "остатки;",
        "ошибки нехватки места или вычислительных ресурсов.",
    ]
)
add_text(
    "Эти данные отображаются в личном кабинете Heritage3D и при загрузке моделей. Если лимитов не хватает, "
    "пользователя можно переводить во встроенные механизмы расширения места или тарифа Geoscan Cloud. "
    "Для корпоративных подписок нужно отдельно обсудить лимиты на уровне команды Geoscan Cloud, "
    "соответствующей музею или исследовательскому центру."
)

doc.add_heading("9. Другие форматы и плееры", level=1)
add_text(
    "Отдельный вопрос — хранение и показ других форматов: PBR-моделей, 2D-сканов, 360-панорам, "
    "гигапиксельных изображений и 360-туров. Поддержка 2D-сканов и панорам 360° рассматривается "
    "в перспективе развития."
)
add_text(
    "Идеально, если все эти медиа можно хранить в Geoscan Cloud, чтобы не переносить файлы на сервер Heritage3D. "
    "Желательно, чтобы встроенный Geoscan viewer открывал такие форматы через Cesium или другой универсальный viewer."
)
add_text(
    "Если это невозможно, то для PlayCanvas, 2D-, 360- и гигапиксельных плееров нужна аналогичная схема:"
)
add_bullets(
    [
        "файлы хранятся в S3-хранилище Geoscan Cloud;",
        "портал получает защищенный manifest или embed;",
        "плеер получает доступ к медиа через временные ссылки или токены;",
        "скачивание исходников можно запрещать;",
        "доступ регулируется правами портала и Geoscan Cloud.",
    ]
)
add_text(
    "Manifest может содержать защищенные ссылки на опубликованные ресурсы в S3 или CDN Geoscan Cloud. "
    "Для приватных моделей, черновиков и материалов с запретом скачивания используются короткоживущие "
    "signed URL или viewer token. Портал не сохраняет прямые ссылки на S3-объекты как постоянные URL."
)

doc.add_heading("10. Embed и manifest для публичного показа", level=1)
add_text(
    "Все плееры моделей должны передаваться на сайт через API в виде Geoscan Cloud embed или manifest "
    "с настраиваемыми параметрами встройки."
)
add_text("Нужны параметры:")
add_bullets(
    [
        "язык;",
        "стартовая камера;",
        "видимость панелей;",
        "включение или выключение туров и POI;",
        "режим доступа;",
        "preview;",
        "размеры iframe.",
    ]
)
add_text("Портал сохраняет у себя:")
add_bullets(
    [
        "geoscanProjectId;",
        "TW_ID;",
        "manifestUrl;",
        "iframeUrl;",
        "previewUrl;",
        "статус публикации;",
        "дату обновления.",
    ]
)

doc.add_heading("11. Авторизация и уровни доступа", level=1)
add_text(
    "Так как авторизация пользователей будет идти через Geoscan ID, доступ к API предлагается строить "
    "через OAuth/OIDC. Пользователь входит через Geoscan ID, после чего портал Heritage3D получает "
    "связанный профиль пользователя и user access token."
)
add_text(
    "Все ручные действия пользователя — загрузка модели, редактирование сцены, создание и изменение POI, "
    "туров, слоев, preview и сохранение — должны выполняться от его имени через geoscan_user_access_token. "
    "Это нужно, чтобы Geoscan Cloud мог проверять команду, роли, права доступа и вести аудит действий."
)
add_text(
    "Отдельно нужен service token для backend Heritage3D. Он используется только для системных операций: "
    "связать H3D_ID, TW_ID и geoscanProjectId; получить статус обработки и preview; сформировать embed; "
    "обработать webhook; синхронизировать технические данные."
)
add_text(
    "Для встроенного редактора нужна короткоживущая editor session token, ограниченная по времени, "
    "пользователю, проекту и разрешенному домену портала."
)
add_text(
    "Для полностью публичных опубликованных моделей embed или viewer может работать без токена через "
    "стабильный публичный iframeUrl или manifestUrl. Для черновиков, модерации, ограниченного доступа, "
    "приватных моделей и сценариев с запретом скачивания нужен короткоживущий embed/viewer token."
)
add_text(
    "Также нужны подпись webhook-уведомлений и проверка происхождения запросов для iframe/postMessage."
)
add_text("Итого получается четыре уровня доступа к API:")
add_numbered(
    [
        "user token — действия пользователя от его имени;",
        "service token — системные операции backend Heritage3D;",
        "editor session token — короткоживущая сессия встроенного редактора;",
        "embed/viewer token — безопасный просмотр черновиков, приватных моделей и объектов с ограниченным доступом.",
    ]
)
add_text("Простой публичный просмотр опубликованной модели возможен без токена.")

doc.add_heading("12. Распределение ответственности", level=1)
add_text("Geoscan Cloud отвечает за:")
add_bullets(
    [
        "S3-хранилище файлов и управление физическими объектами хранения;",
        "обработку и публикацию;",
        "встроенный редактор и viewer;",
        "проекты и слои;",
        "POI, туры и измерения;",
        "preview;",
        "технические статусы;",
        "защищенный доступ к медиаобъектам.",
    ]
)
add_text("Heritage3D отвечает за:")
add_bullets(
    [
        "H3D_ID и TW_ID;",
        "карточку культурного объекта;",
        "Dublin Core и CIDOC CRM;",
        "каталоги и коллекции;",
        "организации и внутренние роли;",
        "модерацию и публикацию;",
        "публичную обвязку;",
        "журнал событий и уведомления.",
    ]
)

doc.add_heading("13. Сопроводительные материалы", level=1)
add_text(
    "Детальный перечень условных эндпойнтов, предполагаемых мест использования на heritage3d.ru, ссылок "
    "на справку Geoscan Cloud и аналогов Sketchfab приведен в отдельной Excel-таблице "
    "«Geoscan Cloud API для Heritage3D.xlsx»."
)

doc.save(OUTPUT)
print(OUTPUT)
